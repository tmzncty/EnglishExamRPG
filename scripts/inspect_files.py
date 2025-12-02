import os
import docx
import pdfplumber

def inspect_docx(file_path):
    print(f"--- Inspecting DOCX: {file_path} ---")
    try:
        doc = docx.Document(file_path)
        # Print first 50 paragraphs to see structure
        for i, para in enumerate(doc.paragraphs[:50]):
            if para.text.strip():
                print(f"Para {i}: {para.text[:100]}")
        
        print("\n... Searching for markers ...")
        markers = ["Section I", "Section II", "Part A", "Part B", "Text 1", "Text 2", "Use of English"]
        for para in doc.paragraphs:
            for marker in markers:
                if marker in para.text:
                    print(f"Found '{marker}': {para.text[:50]}")
    except Exception as e:
        print(f"Error: {e}")

def inspect_pdf(file_path):
    print(f"\n--- Inspecting PDF: {file_path} ---")
    try:
        with pdfplumber.open(file_path) as pdf:
            # Check first 3 pages
            for i, page in enumerate(pdf.pages[:3]):
                text = page.extract_text()
                if text:
                    print(f"Page {i+1} start: {text[:100].replace(chr(10), ' ')}")
            
            print("\n... Searching for markers in first 5 pages ...")
            markers = ["Section I", "Section II", "Part A", "Text 1", "Text 2", "Use of English", "完形填空", "阅读理解"]
            for i, page in enumerate(pdf.pages[:5]):
                text = page.extract_text()
                if text:
                    for marker in markers:
                        if marker in text:
                            print(f"Found '{marker}' on Page {i+1}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(base_dir, "DOCX", "2010年考研英语一真题.docx")
    pdf_path = os.path.join(base_dir, "PDF", "2010年考研英语一真题解析.pdf")
    
    if os.path.exists(docx_path):
        inspect_docx(docx_path)
    else:
        print(f"File not found: {docx_path}")
        
    if os.path.exists(pdf_path):
        inspect_pdf(pdf_path)
    else:
        print(f"File not found: {pdf_path}")
