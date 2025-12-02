"""
从 DOCX 文件中提取图片并转换为 base64
用于图片作文题的显示
"""
import os
import json
import base64
from docx import Document
from pathlib import Path


def extract_images_from_docx(docx_path, output_dir="images"):
    """
    从 DOCX 文件中提取所有图片
    
    Args:
        docx_path: DOCX 文件路径
        output_dir: 输出目录
        
    Returns:
        图片信息列表 [{filename, base64, size}]
    """
    doc = Document(docx_path)
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    images_info = []
    
    # 提取图片
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            
            # 获取文件扩展名
            ext = rel.target_ref.split('.')[-1]
            filename = f"image_{len(images_info) + 1}.{ext}"
            
            # 保存图片文件
            filepath = os.path.join(output_dir, filename)
            with open(filepath, 'wb') as f:
                f.write(image_data)
            
            # 转换为 base64
            base64_data = base64.b64encode(image_data).decode('utf-8')
            
            # 获取 MIME 类型
            mime_type = {
                'png': 'image/png',
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'gif': 'image/gif',
                'bmp': 'image/bmp'
            }.get(ext.lower(), 'image/png')
            
            images_info.append({
                'filename': filename,
                'base64': base64_data,
                'mime_type': mime_type,
                'data_url': f"data:{mime_type};base64,{base64_data}",
                'size': len(image_data)
            })
            
            print(f"✓ 提取图片: {filename} ({len(image_data)} bytes)")
    
    return images_info


def update_json_with_images(json_file, year, images_mapping):
    """
    更新 JSON 文件，添加图片数据
    
    Args:
        json_file: JSON 文件路径
        year: 年份
        images_mapping: {question_id: image_index}
    """
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 提取图片
    docx_file = f"DOCX/{year}年考研英语一真题.docx"
    if not os.path.exists(docx_file):
        print(f"✗ 未找到 DOCX 文件: {docx_file}")
        return
    
    images = extract_images_from_docx(docx_file, f"images/{year}")
    
    # 更新 JSON
    updated = False
    for section in data.get('sections', []):
        for question in section.get('questions', []):
            q_id = question['id']
            if q_id in images_mapping:
                img_idx = images_mapping[q_id]
                if img_idx < len(images):
                    question['image'] = images[img_idx]['data_url']
                    print(f"✓ 为题目 {q_id} 添加图片")
                    updated = True
    
    # 保存更新后的 JSON
    if updated:
        output_file = json_file.replace('.json', '_with_images.json')
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"✓ 已保存到: {output_file}")


def process_2010_exam():
    """处理 2010 年试卷（包含第52题图片）"""
    print("="*60)
    print("处理 2010 年考研英语试卷")
    print("="*60)
    
    # 2010年第52题是图片作文题
    images_mapping = {
        52: 0  # 第52题使用第1张图片（索引0）
    }
    
    update_json_with_images('2010_full.json', 2010, images_mapping)


def process_all_years():
    """处理所有年份的图片作文题"""
    # 定义每年的图片作文题
    image_questions = {
        2010: {52: 0},  # 文化火锅
        2011: {52: 0},  # 旅途之余
        2012: {52: 0},  # 乐观与悲观
        2013: {52: 0},  # 选择
        2014: {52: 0},  # 相携
        2015: {52: 0},  # 手机时代
        2016: {52: 0},  # 教育方式
        2017: {52: 0},  # 读书
        2018: {52: 0},  # 选择
        2019: {52: 0},  # 坚持
        2020: {52: 0},  # 习惯
        2021: {52: 0},  # 教育
        2022: {52: 0},  # 榜样
        2023: {52: 0},  # 奋斗
        2024: {52: 0},  # 继续
    }
    
    for year, mapping in image_questions.items():
        json_file = f"{year}_full.json"
        if os.path.exists(json_file):
            print(f"\n处理 {year} 年...")
            try:
                update_json_with_images(json_file, year, mapping)
            except Exception as e:
                print(f"✗ 处理 {year} 年失败: {str(e)}")


def extract_2010_hotpot_image():
    """专门提取2010年的文化火锅图片"""
    docx_path = "DOCX/2010年考研英语一真题.docx"
    
    if not os.path.exists(docx_path):
        print(f"✗ 文件不存在: {docx_path}")
        return None
    
    print(f"正在从 {docx_path} 提取图片...")
    
    doc = Document(docx_path)
    
    # 查找包含图片的段落
    for i, para in enumerate(doc.paragraphs):
        # 检查段落内容
        if "Part B" in para.text or "52" in para.text:
            print(f"找到 Part B 段落 (段落 {i})")
    
    # 提取所有图片
    images = extract_images_from_docx(docx_path, "images/2010")
    
    if images:
        print(f"\n找到 {len(images)} 张图片:")
        for idx, img in enumerate(images):
            print(f"  [{idx}] {img['filename']} - {img['size']} bytes")
            print(f"      Data URL 前缀: {img['data_url'][:80]}...")
        
        return images
    else:
        print("✗ 未找到图片")
        return None


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        if sys.argv[1] == "2010":
            # 只处理2010年
            images = extract_2010_hotpot_image()
            if images:
                # 更新 JSON
                process_2010_exam()
        elif sys.argv[1] == "all":
            # 处理所有年份
            process_all_years()
    else:
        # 默认处理2010年
        print("提取 2010 年文化火锅图片...\n")
        images = extract_2010_hotpot_image()
        
        if images and len(images) > 0:
            print("\n是否要更新 2010_full.json? (y/n): ", end='')
            choice = input().strip().lower()
            
            if choice == 'y':
                process_2010_exam()
